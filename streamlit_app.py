# =========================
# File: spj_generator.py
# =========================

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def buat_spj(lembaga, nama_kegiatan, tgl, lokasi, anggaran, realisasi, sumber_dana, bukti_upload,
             nama_kades="Sutrisno, S.E.", nama_ketua_bpd="Misdi", nama_ketua_lembaga="", nama_bendahara="", kode_register="470/SPJ-DESA"):

    doc = Document()

    # Tambahkan logo desa jika tersedia
    logo_path = os.path.join(os.path.dirname(__file__), "logo_desa.png")
    if os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.2))

    # KOP DESA
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("PEMERINTAH DESA KELING\n")
    header_run.bold = True
    header_run.font.size = Pt(14)
    header.add_run("KECAMATAN KEPUNG, KABUPATEN KEDIRI\n").bold = True
    header.add_run("Alamat: Jl. Raya Keling, Bukaan, Keling, Kediri, Jawa Timur 64293\n").italic = True
    doc.add_paragraph("_______________________________________________________________")

    # NOMOR REGISTER SURAT
    nomor = f"Nomor: {kode_register}/{tgl.month:02d}/{tgl.year}"
    nomor_paragraf = doc.add_paragraph(nomor)
    nomor_paragraf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    nomor_paragraf.runs[0].bold = True

    # JUDUL
    judul = doc.add_paragraph("SURAT PERTANGGUNGJAWABAN KEGIATAN")
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    judul.runs[0].bold = True
    doc.add_paragraph("\n")

    # ISI DATA KEGIATAN
    doc.add_paragraph(f"Nama Kegiatan        : {nama_kegiatan}")
    doc.add_paragraph(f"Tanggal Pelaksanaan  : {tgl.strftime('%d-%m-%Y')}")
    doc.add_paragraph(f"Lembaga Pelaksana    : {lembaga}")
    doc.add_paragraph(f"Lokasi               : {lokasi}")
    doc.add_paragraph(f"RAB (Anggaran)       : Rp {anggaran:,.0f}")
    doc.add_paragraph(f"Realisasi            : Rp {realisasi:,.0f}")
    doc.add_paragraph(f"Selisih              : Rp {anggaran - realisasi:,.0f}")
    doc.add_paragraph(f"Sumber Dana          : {sumber_dana}")

    doc.add_paragraph("\n")

    # TANGGAL DAN TEMPAT
    ttd_paragraf = doc.add_paragraph()
    ttd_paragraf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd_paragraf.add_run(f"Desa Keling, {tgl.strftime('%d-%m-%Y')}\n")

    # TABEL TANDA TANGAN
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # Baris 1
    table.cell(0, 0).text = "Mengetahui:\nKepala Desa"
    table.cell(0, 1).text = f"Lembaga Pelaksana\nKetua {lembaga}"

    # Baris 2
    table.cell(1, 0).text = nama_kades
    table.cell(1, 1).text = nama_ketua_lembaga if nama_ketua_lembaga else ".................."

    # Baris 3
    table.cell(2, 0).text = ""
    table.cell(2, 1).text = "Bendahara"

    # Baris 4
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = nama_bendahara if nama_bendahara else ".................."

    doc.add_paragraph("\n")

    # TANDA TANGAN KETUA BPD
    bpd_paragraf = doc.add_paragraph()
    bpd_paragraf.add_run("Mengesahkan,\nKetua BPD\n\n\n" + nama_ketua_bpd)
    bpd_paragraf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Simpan dokumen
    out_path = os.path.join(os.path.dirname(__file__), "SPJ_Kegiatan.docx")
    doc.save(out_path)
    return out_path
