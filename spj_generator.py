import os
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import qrcode

def generate_register_code():
    now = datetime.now()
    tanggal = now.strftime("%Y-%m-%d")
    jam = now.strftime("%H%M%S")
    return f"{tanggal}-{jam}"

def remove_table_borders(table):
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)

    tblPr.append(tblBorders)

def buat_spj(lembaga, nama_kegiatan, tgl, lokasi, sumber_dana, rab_df,
             bukti_upload, nama_kades, nama_ketua_bpd, nama_ketua_lembaga,
             nama_bendahara, kode_register):

    doc = Document()

    doc.add_paragraph("PEMERINTAH DESA KELING", style='Heading 1').alignment = 1
    doc.add_paragraph("KECAMATAN KEPUNG, KABUPATEN KEDIRI", style='Heading 2').alignment = 1
    doc.add_paragraph("Jl. Raya Keling, Bukaan, Keling, Kediri, Jawa Timur 64293", style='Normal').alignment = 1
    doc.add_paragraph("______________________________________________________________")
    doc.add_paragraph(f"\nNomor: {kode_register}/{tgl.month:02d}/{tgl.year}")
    doc.add_paragraph("\nSURAT PERTANGGUNGJAWABAN KEGIATAN", style='Title').alignment = 1

    doc.add_paragraph(f"\nNama Kegiatan        : {nama_kegiatan}")
    doc.add_paragraph(f"Tanggal Pelaksanaan  : {tgl.strftime('%d-%m-%Y')}")
    doc.add_paragraph(f"Lembaga Pelaksana    : {lembaga}")
    doc.add_paragraph(f"Lokasi               : {lokasi}")
    doc.add_paragraph(f"Sumber Dana          : {sumber_dana}\n")

    doc.add_paragraph("Tabel RAB dan Realisasi:")
    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["No", "Uraian", "Volume", "Satuan", "Harga Satuan", "Realisasi", "Pajak (%)", "Diskon (Rp)", "Realisasi Bersih"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h

    total_anggaran = 0
    total_realisasi_bersih = 0

    for i, row in rab_df.iterrows():
        try:
            vol = float(row["Volume"])
            harga = float(row["Harga Satuan"])
            realisasi = float(row["Realisasi"])
            pajak_persen = float(row["Pajak (%)"])
            diskon = float(row["Diskon (Rp)"])
        except:
            vol = harga = realisasi = pajak_persen = diskon = 0

        jumlah = vol * harga
        pajak = realisasi * pajak_persen / 100
        realisasi_bersih = realisasi - pajak - diskon

        total_anggaran += jumlah
        total_realisasi_bersih += realisasi_bersih

        data = [
            str(i + 1), str(row["Uraian"]), str(vol), str(row["Satuan"]),
            f"Rp {harga:,.0f}", f"Rp {realisasi:,.0f}",
            f"{pajak_persen:.1f}%", f"Rp {diskon:,.0f}",
            f"Rp {realisasi_bersih:,.0f}"
        ]

        row_cells = table.add_row().cells
        for j, val in enumerate(data):
            row_cells[j].text = val

    doc.add_paragraph(f"\nTotal Anggaran        : Rp {total_anggaran:,.0f}")
    doc.add_paragraph(f"Total Realisasi Bersih: Rp {total_realisasi_bersih:,.0f}")
    doc.add_paragraph(f"Selisih               : Rp {total_anggaran - total_realisasi_bersih:,.0f}\n")

    # Tanda tangan
    ttd_table = doc.add_table(rows=3, cols=2)
    remove_table_borders(ttd_table)
    ttd_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    ttd_table.cell(0, 0).text = f"Mengetahui,\nKepala Desa\n\n\n{nama_kades}"
    ttd_table.cell(0, 1).text = f"Lembaga Pelaksana\nKetua {lembaga}\n\n\n{nama_ketua_lembaga}"
    ttd_table.cell(2, 0).text = f"Bendahara\n\n\n{nama_bendahara}"
    ttd_table.cell(2, 1).text = f"Mengesahkan,\nKetua BPD\n\n\n{nama_ketua_bpd}"

    doc.add_paragraph("\n")

    # QR Code
    qr_data = f"Dokumen SPJ {kode_register} - Desa Keling"
    qr = qrcode.make(qr_data)
    qr_path = "qr_temp.png"
    qr.save(qr_path)
    doc.add_picture(qr_path, width=Inches(1.2))
    os.remove(qr_path)

    doc.add_paragraph(
        "Barcode ini menunjukkan dokumen resmi yang diterbitkan oleh Pemerintah Desa Keling, Kec. Kepung, Kab. Kediri.",
        style="Normal"
    ).alignment = 1

    output_path = "SPJ_Kegiatan.docx"
    doc.save(output_path)
    return output_path
